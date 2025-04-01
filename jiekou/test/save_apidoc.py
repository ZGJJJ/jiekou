# save_api_doc.py
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


def create_api_documentation():
    # 创建文档
    doc = Document()

    # 设置中文字体
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 添加标题
    title = doc.add_heading('API 接口文档2', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基础信息
    doc.add_heading('基础信息', 1)
    doc.add_paragraph('基础URL: http://localhost:5000')
    p = doc.add_paragraph('所有需要认证的接口都需要在请求头中包含：')
    doc.add_paragraph('• Authorization: Bearer <access_token> (访问令牌)')
    doc.add_paragraph('• Refresh-Token: <refresh_token> (刷新令牌，用于自动刷新访问令牌)')
    doc.add_paragraph('• X-API-Key: <api_key> (API密钥)')

    # 认证机制
    doc.add_heading('认证机制', 1)
    doc.add_paragraph('• 访问令牌(access_token)有效期：60分钟')
    doc.add_paragraph('• 刷新令牌(refresh_token)有效期：90天')
    doc.add_paragraph('• 当访问令牌过期时，系统会自动使用刷新令牌获取新的令牌对')

    # 用户注册接口
    doc.add_heading('1. 用户注册', 1)
    doc.add_paragraph('接口: /register')
    doc.add_paragraph('方法: POST')
    doc.add_paragraph('描述: 注册新用户并获取API密钥')

    doc.add_heading('请求体:', 2)
    doc.add_paragraph('''
{
    "username": "用户名",
    "password": "密码"
}
''')

    doc.add_heading('响应示例:', 2)
    doc.add_paragraph('''
{
    "message": "Registration successful",
    "api_key": "生成的API密钥"
}
''')

    # 用户登录接口
    doc.add_heading('2. 用户登录', 1)
    doc.add_paragraph('接口: /login')
    doc.add_paragraph('方法: POST')
    doc.add_paragraph('描述: 用户登录并获取令牌')

    doc.add_heading('请求体:', 2)
    doc.add_paragraph('''
{
    "username": "用户名",
    "password": "密码"
}
''')

    doc.add_heading('响应示例:', 2)
    doc.add_paragraph('''
{
    "access_token": "访问令牌",
    "refresh_token": "刷新令牌",
    "api_key": "API密钥"
}
''')

    # 企业信息查询接口
    doc.add_heading('3. 企业信息查询', 1)
    doc.add_paragraph('接口: /query')
    doc.add_paragraph('方法: POST')
    doc.add_paragraph('描述: 查询企业相关信息')
    doc.add_paragraph('需要认证: 是')

    doc.add_heading('请求头:', 2)
    doc.add_paragraph('• Authorization: Bearer <access_token>')
    doc.add_paragraph('• Refresh-Token: <refresh_token>')
    doc.add_paragraph('• X-API-Key: <api_key>')

    doc.add_heading('请求体:', 2)
    doc.add_paragraph('''
{
    "company_name": "企业名称"
}
''')

    doc.add_heading('响应示例:', 2)
    doc.add_paragraph('''
{
    "data": [
        {
            "ename": "企业名称",
            "credit_code": "统一社会信用代码",
            "score": "总分",
            "rating": "评级",
            "business_score": "业务得分",
            "undertake_score": "承接能力得分",
            "stability_score": "稳定性得分",
            "performance_score": "履约得分",
            "risk_score": "风险得分",
            "performance_appraisal": "履约评价",
            "bad_behavior_3y": "近3年不良行为",
            "malicious_events_1y": "近1年恶意事件",
            "is_blacklist": "是否黑名单"
        }
    ]
}
''')

    # API使用统计接口
    doc.add_heading('4. API使用统计', 1)
    doc.add_paragraph('接口: /usage')
    doc.add_paragraph('方法: GET')
    doc.add_paragraph('描述: 查询API调用使用统计')
    doc.add_paragraph('需要认证: 是')

    doc.add_heading('请求头:', 2)
    doc.add_paragraph('• Authorization: Bearer <access_token>')
    doc.add_paragraph('• Refresh-Token: <refresh_token>')
    doc.add_paragraph('• X-API-Key: <api_key>')

    doc.add_heading('查询参数:', 2)
    doc.add_paragraph('• group_by: 统计维度（可选值：hour/day/month/year，默认：day）')
    doc.add_paragraph('• start_date: 开始日期（可选，格式：YYYY-MM-DD）')
    doc.add_paragraph('• end_date: 结束日期（可选，格式：YYYY-MM-DD）')

    doc.add_heading('响应示例:', 2)
    doc.add_paragraph('''
{
    "usage": [
        {
            "endpoint": "API端点名称",
            "year": 2024,
            "month": 3,
            "day": 20,
            "hour": 14,          // 仅在 group_by=hour 时显示
            "success_count": 10,  // 成功调用次数
            "fail_count": 2,      // 失败调用次数
            "total_data_count": 50, // 返回数据总条数
            "total_count": 12      // API调用总次数
        }
    ],
    "group_by": "hour",
    "start_date": "2024-03-20",
    "end_date": "2024-03-20"
}
''')

    # 错误响应
    doc.add_heading('错误响应', 1)
    doc.add_paragraph('所有接口的错误响应格式统一为：')
    doc.add_paragraph('''
{
    "error": "错误信息",
    "error_code": "错误代码",
    "message": "用户友好提示信息"
}
''')

    doc.add_heading('常见错误代码:', 2)
    doc.add_paragraph('• TOKEN_MISSING: 缺少JWT令牌')
    doc.add_paragraph('• TOKEN_INVALID: 无效的令牌')
    doc.add_paragraph('• ACCESS_TOKEN_EXPIRED: 访问令牌过期')
    doc.add_paragraph('• REFRESH_TOKEN_EXPIRED: 刷新令牌过期')
    doc.add_paragraph('• REFRESH_TOKEN_INVALID: 无效的刷新令牌')

    # 使用示例
    doc.add_heading('使用示例', 1)
    doc.add_heading('cURL示例:', 2)

    doc.add_paragraph('1. 注册新用户：')
    doc.add_paragraph('''
curl -X POST http://localhost:5000/register \\
     -H "Content-Type: application/json" \\
     -d '{"username": "testuser", "password": "testpass"}'
''')

    doc.add_paragraph('2. 用户登录：')
    doc.add_paragraph('''
curl -X POST http://localhost:5000/login \\
     -H "Content-Type: application/json" \\
     -d '{"username": "testuser", "password": "testpass"}'
''')

    doc.add_paragraph('3. 查询企业信息：')
    doc.add_paragraph('''
curl -X POST http://localhost:5000/query \\
     -H "Authorization: Bearer <access_token>" \\
     -H "Refresh-Token: <refresh_token>" \\
     -H "X-API-Key: <api_key>" \\
     -H "Content-Type: application/json" \\
     -d '{"company_name": "ABB(中国)有限公司"}'
''')

    doc.add_paragraph('4. 查询使用统计：')
    doc.add_paragraph('''
curl -X GET "http://localhost:5000/usage?group_by=hour&start_date=2024-03-20" \\
     -H "Authorization: Bearer <access_token>" \\
     -H "Refresh-Token: <refresh_token>" \\
     -H "X-API-Key: <api_key>"
''')

    # 注意事项
    doc.add_heading('注意事项', 1)
    doc.add_paragraph('1. 所有日期格式必须为：YYYY-MM-DD')
    doc.add_paragraph('2. 访问令牌过期会自动使用刷新令牌获取新令牌')
    doc.add_paragraph('3. 刷新令牌过期需要重新登录')
    doc.add_paragraph('4. API调用统计包含成功/失败次数和数据条数')
    doc.add_paragraph('5. 统计查询支持多种时间维度的聚合')

    # 保存文档
    doc.save('API接口文档.docx')


if __name__ == '__main__':
    create_api_documentation()